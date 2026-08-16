from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "SaQshi_Health_to_Education_Change_Reference.docx"
LOGO = ROOT / "ui" / "assets" / "images" / "bihar-education-department-logo.jpeg"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
PALE = "F4F6F9"
GRAY = "5B6573"
WHITE = "FFFFFF"
RED = "DC2626"
AMBER = "F59E0B"
GREEN = "16A34A"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        cell.text = text
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(NAVY)
                run.font.size = Pt(9.5)
    for values in rows:
        row = table.add_row()
        for i, value in enumerate(values):
            cell = row.cells[i]
            cell.text = str(value)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.size = Pt(9.25)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def add_note(doc, title, text, fill=PALE):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "D8E1EF")
        borders.append(el)
    p_pr.append(borders)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title + " ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text)


def configure_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    settings = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in settings.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = "SaQshi Education | Change Reference"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in p.runs:
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor.from_string(GRAY)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "Internal reference | Prepared 06 August 2026 | tech4gov@piramalswasthya.org"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in fp.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor.from_string(GRAY)


def build():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)

    # Editorial cover pattern.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    if LOGO.exists():
        shape = p.add_run().add_picture(str(LOGO), width=Inches(1.45))
        doc_pr = shape._inline.docPr
        doc_pr.set("descr", "Education Department, Government of Bihar logo")
        doc_pr.set("title", "Bihar Education Department")
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(12)
    kicker.paragraph_format.space_after = Pt(8)
    kr = kicker.add_run("CHANGE REFERENCE")
    kr.bold = True
    kr.font.size = Pt(10)
    kr.font.color.rgb = RGBColor.from_string(BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    tr = title.add_run("SaQshi Health to SaQshi Education")
    tr.bold = True
    tr.font.size = Pt(27)
    tr.font.color.rgb = RGBColor.from_string(NAVY)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(26)
    sr = sub.add_run("Functional, UI, Data, Analytics and Deployment Changes")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor.from_string(GRAY)
    add_note(doc, "Purpose:", "A concise reference for stakeholders, demonstrations, testing and deployment teams describing how the original health-oriented SaQshi workflow was adapted for school education.")
    add_table(doc, ["Document", "Value"], [
        ("Build context", "SaQshi Education / School Quality Assessment"),
        ("Prepared on", "06 August 2026"),
        ("Primary audience", "Product, programme, implementation, QA and operations teams"),
        ("Support", "tech4gov@piramalswasthya.org"),
    ], [2700, 6660])
    doc.add_page_break()

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph("The solution was converted from a health-facility assessment platform into an education quality assessment platform while retaining the reusable SaQshi assessment engine. The conversion is driven mainly through an Education domain profile, an Education scoring framework and education-specific screens, analytics and labels.")
    add_bullet(doc, "Schools are the assessed units; UDISE Code is the school identifier.")
    add_bullet(doc, "Classes, Domains and Questions replace health-oriented Departments, Areas of Concern and Checkpoints in user-facing workflows.")
    add_bullet(doc, "Education categorisation uses Abhilasha, Pragati and Jagriti score bands.")
    add_bullet(doc, "State dashboards, maps, drill-downs and indicator analytics now present education data at school, block, district and division levels.")
    add_bullet(doc, "Assessor information captures classroom/teacher details with education-specific validation.")
    add_bullet(doc, "Branding, login, footer contact and deployment hardening were updated for the Education implementation.")

    doc.add_heading("2. Core terminology conversion", level=1)
    add_table(doc, ["Health-oriented term", "Education term", "Where used"], [
        ("Facility", "School", "Dashboards, mappings, drill-down, reports and maps"),
        ("NIN / facility code", "UDISE Code", "School identity and search"),
        ("Department", "Class", "Assessment selection and checklist entry"),
        ("Area of Concern", "Domain", "Scoring, analytics and completion messages"),
        ("Checkpoint", "Question", "Questionnaire and general navigation"),
        ("Facility profile", "School Profile", "Administration and drill-down"),
        ("Facility drill-down", "School Drill-down", "State monitoring"),
        ("Mapped facilities", "Mapped Schools", "Assessor management and map view"),
        ("Assessment checklist", "Questionnaire", "Education assessment workflow"),
    ], [2300, 2200, 4860])
    add_note(doc, "Configuration approach:", "The Education vocabulary and enabled modules are defined in api/config/domain.json and api/config/profiles/education.json. This keeps the base product reusable and reduces hard-coded health terminology.")

    doc.add_heading("3. Branding and login experience", level=1)
    add_bullet(doc, "Login messaging was changed to school quality and structured school improvement language.")
    add_bullet(doc, "The Bihar Education Department logo is shown with the SaQshi logo on the login screen.")
    add_bullet(doc, "The public branding configuration supports separate primary and secondary login logos and accessible alternative text.")
    add_bullet(doc, "Footer support email was standardized to tech4gov@piramalswasthya.org.")
    add_bullet(doc, "Asset query versions and service-worker cache versions were updated so users receive the changed login and dashboard assets.")

    doc.add_heading("4. Education assessment framework", level=1)
    doc.add_paragraph("The default framework is saqshi-education. A class is the assessment unit and school performance is aggregated from completed class scores in the assessment round. The configured domains are:")
    for d in ["Teaching Planning and Preparedness", "Students' Engagement", "Assessment and Feedback", "Classroom Learning Environment & Culture", "Learning Closure and Reflection"]:
        add_bullet(doc, d)
    doc.add_heading("Performance categories", level=2)
    table = add_table(doc, ["Category", "Percentage rule", "Map/dashboard colour", "Meaning in visualisation"], [
        ("Abhilasha", "Below 60%", "Red", "Needs the greatest improvement"),
        ("Pragati", "60% to 75%, inclusive", "Amber", "Developing / progressing"),
        ("Jagriti", "Above 75%", "Green", "Stronger performance"),
    ], [1800, 2600, 2200, 2760])
    for i, color in enumerate((RED, AMBER, GREEN), start=1):
        set_cell_shading(table.rows[i].cells[2], color)
        for run in table.rows[i].cells[2].paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.bold = True

    doc.add_heading("5. Assessor workflow changes", level=1)
    add_table(doc, ["Area", "Education implementation"], [
        ("Assessor details", "Captures assessor identity together with class/teacher information."),
        ("Teacher name", "Displayed as Class / Subject Teacher's Name."),
        ("Teacher identifier", "Displayed as Teacher ID and accepts letters and digits only; special characters are rejected."),
        ("Subject", "Separate Subject field for the education workflow."),
        ("Assessment Type", "Removed from the visible form; the required internal value remains handled by the application."),
        ("Name validation", "Teacher/person name accepts alphabetic characters and spaces, with matching frontend and backend checks."),
        ("Checklist prompts", "Select class to begin checklist; completion messaging refers to the Domain."),
        ("Persistence", "Teacher ID is stored in teacher_code and Subject in subject_name."),
    ], [2200, 7160])

    doc.add_heading("6. State dashboard and categorisation", level=1)
    add_bullet(doc, "Colored summary cards show the number of schools in Abhilasha, Pragati and Jagriti.")
    add_bullet(doc, "Expandable Division and District sections show each area's percentage, resulting category and the distribution of assessed schools across all three categories.")
    add_bullet(doc, "Division-wise and district-wise bar graphs use the same category colors and provide hover details.")
    add_bullet(doc, "Categorisation panels and graphs can be downloaded as PNG images for presentation and reporting.")
    add_bullet(doc, "Backend aggregation uses the latest completed assessment data and is exposed through the state dashboard response.")

    doc.add_heading("7. School map changes", level=1)
    add_table(doc, ["Capability", "Change"], [
        ("Map size", "Map View card and map canvas were enlarged for clearer state-level viewing."),
        ("Mapped-school list", "Displayed below the map with five schools per page to preserve map space."),
        ("Plot selector", "Users can switch between Facility plotting and Domain-wise categorisation."),
        ("Facility plotting", "Plots individual school GPS points."),
        ("Domain plotting", "Colors district polygons by overall or selected-domain percentage and hides facility points."),
        ("Hover information", "Shows district percentage/category for the selected overall/domain view."),
        ("Legend", "Displays Abhilasha, Pragati, Jagriti and No score colors below/within the map view."),
        ("Point clarity", "School markers were reduced to small points with a lighter border to reduce overlap."),
        ("District matching", "District-name aliases/normalisation are used where map JSON and database naming differ."),
    ], [2200, 7160])

    doc.add_heading("8. Drill-down and analytics changes", level=1)
    doc.add_heading("School Drill-down", level=2)
    add_bullet(doc, "Hierarchy follows State > Division > District > Block > School.")
    add_bullet(doc, "The header uses a loading state instead of briefly displaying 0 Schools before hierarchy data arrives.")
    add_bullet(doc, "Selection instructions refer to a school, and domain-wise score details are available for the selected school.")
    add_bullet(doc, "Hierarchy loading is resilient when an optional score aggregation is unavailable.")
    doc.add_heading("Indicator and Field Analytics", level=2)
    add_bullet(doc, "Indicator Analytics wording now ranks checklist indicators by schools in the lowest configured score band.")
    add_bullet(doc, "The title reads Checklist Indicators With Most Low-Score Schools.")
    add_bullet(doc, "Class is included where needed in indicator results.")
    add_bullet(doc, "Actual checkpoint/question text is displayed instead of internal identifiers such as Checkpoint 10001.")
    add_bullet(doc, "Health-specific Department and Standard fields were removed from Field Analytics output.")
    add_bullet(doc, "Low-score school detail can be downloaded for further analysis.")

    doc.add_heading("9. Administration changes", level=1)
    add_table(doc, ["Module", "Change"], [
        ("Assessor Management", "Education labels use schools and UDISE. Selecting an assessor shows associated schools with pagination of 10 records per page, Previous/Next controls and a visible range/total."),
        ("School mapping", "Search, selection and bulk assignment support mapping schools to an assessor."),
        ("User Administration", "Profile update supports first, middle and last name, username, email, mobile and optional password. Access is restricted by role; unsupported edit actions are removed."),
        ("Profile security", "Sensitive profile values continue to use the application's encrypted storage conventions."),
    ], [2200, 7160])

    doc.add_heading("10. Backend, schema and data changes", level=1)
    add_table(doc, ["Change area", "Reference"], [
        ("Education profile", "api/config/profiles/education.json"),
        ("Active domain labels/branding", "api/config/domain.json"),
        ("Education scoring", "api/config/scoring/saqshi-education.json"),
        ("Teacher Subject migration", "api/sql/2026_08_06_assessor_info_education_fields.sql"),
        ("Teacher ID migration", "api/sql/2026_08_06_assessor_info_teacher_code.sql"),
        ("Base schema", "assessment_assessor_info includes teacher_code and subject_name"),
        ("Demo assessment data", "api/sql/demo/2026_08_06_school_assessment_demo_seed.sql"),
        ("Map category demo data", "api/sql/demo/2026_08_06_school_map_category_demo_seed.sql"),
        ("Hierarchy demo data/indexes", "Hierarchy category seed and state hierarchy score indexes under api/sql"),
    ], [2600, 6760])
    add_note(doc, "Demo data intent:", "The demo seed deliberately mixes Abhilasha, Pragati and Jagriti across divisions and districts so dashboard cards, graphs, maps and drill-down screens are understandable during demonstrations. It is test/demo data and should not be represented as production assessment evidence.")

    doc.add_heading("11. Reliability and production hardening", level=1)
    add_bullet(doc, "State dashboard sections use safe defaults so one optional aggregation does not cause the complete dashboard response to fail.")
    add_bullet(doc, "School hierarchy remains usable when optional score calculations fail, while core hierarchy data is still returned.")
    add_bullet(doc, "Login cryptography validates the configured RSA key and returns a controlled authentication error instead of exposing an unhandled server failure.")
    add_bullet(doc, "Frontend asset cache-busting was applied to changed pages to reduce stale browser/service-worker content.")
    add_note(doc, "Production note:", "A browser request to http://localhost on a production client indicates that the deployed API base URL or environment configuration is still pointing to localhost. Set the production API base path/host, confirm database and encryption-key configuration, and then clear the application cache/service worker.", fill="FFF4E5")

    doc.add_heading("12. Deployment checklist", level=1)
    checklist = [
        "Back up the production database and application configuration.",
        "Deploy the Education profile, scoring framework, UI assets, backend services and SQL migration files.",
        "Run the Teacher Subject and Teacher ID migrations once in the target database.",
        "Confirm domain.json selects the education profile and saqshi-education framework.",
        "Confirm the Bihar Education Department logo and SaQshi logo load from the deployed public path.",
        "Set the production API base URL; verify no production browser request targets localhost.",
        "Verify RSA/encryption configuration and writable application log paths.",
        "Clear or advance service-worker/static-asset caches after deployment.",
        "Test login for State Admin, Assessor and relevant school/user roles.",
        "Test dashboard category cards, division/district expansion, graphs and PNG downloads.",
        "Test Facility and Domain map modes, district hover values, legend and mapped-school pagination.",
        "Test School Drill-down hierarchy, Indicator Analytics and low-score downloads.",
        "Test assessor information validation and persistence of Teacher ID and Subject.",
        "Keep demo seed scripts out of production unless explicitly approved for a demonstration environment.",
    ]
    for item in checklist:
        add_bullet(doc, "[ ] " + item)

    doc.add_heading("13. Acceptance reference", level=1)
    add_table(doc, ["Test", "Expected result"], [
        ("Education terminology", "No health-facility language appears in the primary Education workflows."),
        ("Category boundary", "59.99% is Abhilasha; 60–75% is Pragati; above 75% is Jagriti."),
        ("Map mode", "Facility mode shows small school points; Domain mode shows district polygons without facility points."),
        ("Teacher ID", "Letters and digits are accepted; spaces and special characters are rejected."),
        ("Assessment Type", "Not visible in the Education assessor-information form."),
        ("Assessor mappings", "More than 10 mapped schools produce paginated pages."),
        ("School hierarchy", "Shows Loading schools while fetching and then displays the returned school total."),
        ("Footer", "Support contact is tech4gov@piramalswasthya.org."),
    ], [2600, 6760])

    doc.add_heading("14. Scope note", level=1)
    doc.add_paragraph("This is a working change reference, not a formal release note or exhaustive code diff. It records the principal functional and technical adaptations present in the Education build and the deployment actions needed to use them safely. Exact record counts vary by environment and assessment data.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "SaQshi Health to SaQshi Education - Change Reference"
    doc.core_properties.subject = "Education conversion reference"
    doc.core_properties.author = "SaQshi Project Team"
    doc.core_properties.keywords = "SaQshi, education, school assessment, change reference"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
