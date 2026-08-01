from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "docs" / "user" / "SaQshi_DPO_Assessor_One_Page_Guide_v2.docx"

BLUE = RGBColor(31, 78, 120)
DARK = RGBColor(11, 37, 69)
LIGHT = "E8EEF5"

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=60, start=100, bottom=60, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def font(run, size=9.5, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def p(doc, text="", size=9.5, bold=False, color=None, after=2, before=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.05
    font(para.add_run(text), size, bold, color)
    return para

def heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(5)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.keep_with_next = True
    font(para.add_run(text), 10.5, True, BLUE)

def numbered(doc, items):
    for index, text in enumerate(items, 1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.first_line_indent = Inches(-0.15)
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.line_spacing = 1.03
        font(para.add_run(f"{index}. "), 9.2, True, DARK)
        font(para.add_run(text), 9.2)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.55)
sec.bottom_margin = Inches(0.55)
sec.left_margin = Inches(0.60)
sec.right_margin = Inches(0.60)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(9.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(1)
font(title.add_run("SaQshi DPO / Assessor Assessment Guide"), 17, True, DARK)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(5)
font(sub.add_run("State Admin setup, assignment scope and checklist workflow"), 9.5, False, BLUE)

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
headers = ["Role", "Scope", "Key responsibility"]
for i, text in enumerate(headers):
    cell = table.rows[0].cells[i]
    shade(cell, LIGHT); set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    font(cell.paragraphs[0].add_run(text), 8.8, True, DARK)
for role, scope, task in [
    ("State Admin", "State-wide", "Create DPO/Assessor accounts; assign, end or transfer School/Facility mappings; monitor reports."),
    ("DPO / Assessor", "Mapped units only", "Start or resume assessment, complete checklist and download assessment report."),
    ("School / Facility User", "Own unit", "Uses the normal facility workflow and views own data.")]:
    cells = table.add_row().cells
    for cell, text in zip(cells, (role, scope, task)):
        set_cell_margins(cell)
        cell.text = ""
        font(cell.paragraphs[0].add_run(text), 8.5, cell == cells[0], DARK if cell == cells[0] else None)

heading(doc, "1. State Admin: create and assign DPO / Assessor")
numbered(doc, [
    "Open State Monitoring > DPO / Assessor Management. Enter code, name, designation, mobile and email; leave Linked User ID blank to create login automatically.",
    "Click Save. The code becomes the username. Copy the temporary password shown on screen and share it securely.",
    "Select the DPO/Assessor, search School/Facility by name, UDISE/NIN, district or block, then select one or more units and click Assign selected."
])

heading(doc, "2. Assignment scope and transfer")
p(doc, "A DPO/Assessor can assess only active mapped Schools/Facilities. Education uses School and UDISE Code; Healthcare uses Facility and NIN.", 9.2)
p(doc, "To transfer a unit: select the current DPO/Assessor > click End Assignment > confirm > select the next DPO/Assessor > assign the same unit. Old mapping and assessment history remain available.", 9.2, bold=True, color=DARK)

heading(doc, "3. DPO / Assessor: complete checklist")
numbered(doc, [
    "Login, change the temporary password if asked, and open DPO / Assessor Dashboard.",
    "Use Start Assessment for a new unit or Start Reassessment for a completed unit. Activate required Department/Section.",
    "Save DPO/Assessor and Assessee details. Open Checklist, select Department/Section, Area of Concern and Standard.",
    "Answer and save every checkpoint; attach remarks/evidence where required. SaQshi automatically closes the Department/Section and then the assessment when all active checkpoints are complete."
])

heading(doc, "4. Completion, reassessment and reports")
p(doc, "Completed assessment: dashboard shows Start Reassessment and creates name/dates automatically. Unfinished assessment: use Cancel; saved responses remain in history. Reports show score, checkpoint totals and Score Trend. Search trends by School/Facility name or UDISE/NIN.", 9.2)

flow = doc.add_paragraph()
flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
flow.paragraph_format.space_before = Pt(5)
flow.paragraph_format.space_after = Pt(0)
font(flow.add_run("Create DPO/Assessor  ->  Assign unit  ->  Start assessment  ->  Checklist complete  ->  Completed  ->  Reassessment when due"), 8.7, True, BLUE)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "SaQshi DPO Assessor One-Page Guide"
doc.core_properties.author = "SaQshi"
doc.save(OUT)
print(OUT)
